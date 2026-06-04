import io
import docx
import pdfplumber
import logging
from typing import Union

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_contents: bytes) -> str:
    """Extract plain text from PDF binary contents using pdfplumber"""
    text_content = []
    try:
        with pdfplumber.open(io.BytesIO(file_contents)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                else:
                    logger.warning(f"Could not extract text from PDF page {i+1}")
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF file: {e}")

def extract_text_from_docx(file_contents: bytes) -> str:
    """Extract plain text from Word (.docx) binary contents"""
    try:
        doc = docx.Document(io.BytesIO(file_contents))
        paragraphs_text = [p.text for p in doc.paragraphs]
        # Also extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_cells))
        
        full_text = paragraphs_text + table_text
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX file: {e}")

def extract_text_from_bytes(file_contents: bytes, filename: str) -> str:
    """Detect file type by extension and extract plain text content"""
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_contents)
    elif fn_lower.endswith(".docx"):
        return extract_text_from_docx(file_contents)
    elif fn_lower.endswith((".txt", ".md", ".json")):
        try:
            return file_contents.decode("utf-8")
        except UnicodeDecodeError:
            return file_contents.decode("latin-1")
    else:
        raise ValueError("Unsupported file format. Only PDF, DOCX, and text files are supported.")
